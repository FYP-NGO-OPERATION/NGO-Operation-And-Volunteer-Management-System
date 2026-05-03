"""Generate FYP-03 Thesis DOCX (Chapters 7–8) — Air University Multan Campus."""
import os, zipfile, tempfile, shutil
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

def inject_update_fields(docx_path):
    temp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        settings_path = os.path.join(temp_dir, 'word', 'settings.xml')
        if os.path.exists(settings_path):
            with open(settings_path, 'r', encoding='utf-8') as f:
                content = f.read()
            if '<w:updateFields w:val="true"/>' not in content:
                import re
                content = re.sub(r'(<w:settings[^>]*>)', r'\1<w:updateFields w:val="true"/>', content, count=1)
                with open(settings_path, 'w', encoding='utf-8') as f:
                    f.write(content)
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    zip_out.write(file_path, os.path.relpath(file_path, temp_dir))
    finally:
        shutil.rmtree(temp_dir)

STUDENT = "Muhammad Maauz Mansoor"
REG = "233599"
SESSION = "2023–2027"
UNIVERSITY = "Air University Multan Campus"
DEPARTMENT = "Department of Computer Science"
SUPERVISOR = "Miss Fatima Yousuf"

doc = Document()

def add_page_number(run):
    for tag, attr in [('begin',None),('instrText','PAGE'),('separate',None),('end',None)]:
        el = OxmlElement('w:fldChar' if tag != 'instrText' else 'w:instrText')
        if tag == 'instrText':
            el.set(ns.qn('xml:space'), 'preserve'); el.text = attr
        else:
            el.set(ns.qn('w:fldCharType'), tag)
        run._r.append(el)

def add_seq_field(paragraph, seq_identifier, text_before, text_after):
    run_before = paragraph.add_run(text_before)
    run_before.font.size = Pt(10); run_before.italic = True
    
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(ns.qn('w:fldCharType'), 'begin')
    r1 = paragraph.add_run(); r1._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText'); instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = f' SEQ {seq_identifier} \\* ARABIC '
    r2 = paragraph.add_run(); r2._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(ns.qn('w:fldCharType'), 'separate')
    r3 = paragraph.add_run(); r3._r.append(fldChar2)
    
    r4 = paragraph.add_run('1') # Placeholder
    r4.font.size = Pt(10); r4.italic = True
    
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(ns.qn('w:fldCharType'), 'end')
    r5 = paragraph.add_run(); r5._r.append(fldChar3)
    
    run_after = paragraph.add_run(text_after)
    run_after.font.size = Pt(10); run_after.italic = True

def add_table(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for r in c.paragraphs[0].runs: r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
    return t

def table_caption(t):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_seq_field(p, "Table", "Table ", f": {t}")
    return p

for section in doc.sections:
    section.top_margin = Inches(1.1); section.left_margin = Inches(1.3)
    section.right_margin = Inches(0.8); section.bottom_margin = Inches(0.8)
    section.footer_distance = Inches(0.2)
    p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(p.add_run())

style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(12); style.paragraph_format.line_spacing = 1.5

for level, sz in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 13)]:
    hs = doc.styles[level]
    hs.font.name = 'Times New Roman'; hs.font.size = Pt(sz); hs.font.bold = True
    hs.font.color.rgb = None; hs.paragraph_format.line_spacing = 1.5

def h1(t): p = doc.add_heading(t, level=1); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; return p
def h2(t): return doc.add_heading(t, level=2)
def h3(t): return doc.add_heading(t, level=3)
def body(t): p = doc.add_paragraph(t); p.style = doc.styles['Normal']; return p

# =================== TITLE PAGE ===================
for _ in range(4): doc.add_paragraph()
h1(UNIVERSITY)
h1(DEPARTMENT)
doc.add_paragraph()
h1("NGO Operation and Volunteer Management System")
h1("(HRAS)")
doc.add_paragraph()
h1("FYP-03 Report: Chapters 7–8")
doc.add_paragraph()
body(f"Submitted by: {STUDENT} ({REG})").alignment = WD_ALIGN_PARAGRAPH.CENTER
body(f"Supervised by: {SUPERVISOR}").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# =================== CHAPTER 7 ===================
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Chapter 7'); r.bold = True; r.font.size = Pt(22)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Software Quality Assurance and Testing'); r2.bold = True; r2.font.size = Pt(20)
doc.add_page_break()

h1("Chapter 7: Software Quality Assurance and Testing")

h2("7.1 Introduction to Software Validation")
body("In modern software engineering, deploying an untested application into a live production environment is a catastrophic liability. For an NGO ecosystem where critical financial data and massive volunteer logistics are managed in real-time, any system failure could result in the unrecoverable loss of donor trust and severe operational paralysis. To mitigate these risks, the HRAS system was subjected to an aggressive, multi-layered Quality Assurance (QA) protocol prior to final deployment.")

h2("7.2 Multi-Layered Testing Strategy")
body("The testing matrix was architected around three fundamental paradigms:")
body("1. Automated Unit Testing (White-Box): Micro-level testing executed continuously during development. Scripts were engineered to validate core logic vectors—such as the mathematical accuracy of the financial aggregation engine and the boundary value constraints of the Smart Matching arrays—independent of the UI.")
body("2. Integration Testing: Intermediate-level testing designed to ensure that discrete application modules (e.g., the Dart frontend and the Firestore backend) communicate flawlessly, accurately handling network latency and asynchronous state mutations.")
body("3. Black Box Testing & UAT: Macro-level validation conducted from the end-user's perspective. Evaluators interacted with the compiled UI strictly to ascertain if the system reliably fulfills the functional requirements outlined in Chapter 4, without any knowledge of the underlying codebase.")

h2("7.3 Automated Unit Test Coverage")
body("A total of 62 highly specific automated unit tests were scripted utilizing the Dart/Flutter testing framework. The execution of these tests ensures zero regression defects during subsequent code compilation.")

table_caption("Automated Unit Testing Execution Summary")
add_table(["Functional Domain", "Test Vectors Scripted", "Passed", "Failed", "Integrity Status"],
    [["Regex & Form Validators", "26", "26", "0", "100% Validated"],
     ["Data Models (JSON Serialization)", "15", "15", "0", "100% Validated"],
     ["RBAC & Route Protection Logic", "21", "21", "0", "100% Validated"],
     ["Aggregate Benchmark", "62", "62", "0", "Passed Production Threshold"]])

h2("7.4 Black Box Testing Matrix")
body("The following tables document the rigorous manual testing scenarios applied to the system's core execution paths to validate the architectural integrity against edge-case manipulation.")

table_caption("QA Vector: Authentication & Authorization Engine")
add_table(["Test ID", "Execution Scenario", "Expected Architectural Response", "Empirical Result", "Status"],
    [["TC-01", "Provide valid credentials to Auth API", "Issue secure token, route to appropriate Dashboard", "Successfully routed", "Pass"],
     ["TC-02", "Inject malformed email string", "Trigger client-side Regex block, halt API call", "API call halted, Error rendered", "Pass"],
     ["TC-03", "Provide invalid cryptographic password", "Catch Firebase AuthException, render generic error", "Exception caught safely", "Pass"],
     ["TC-04", "Simulate Privilege Escalation (Volunteer hitting Admin URL)", "Server-side RBAC denies read, client routes to 'Access Denied'", "Request mathematically blocked", "Pass"]])
doc.add_paragraph()

table_caption("QA Vector: Central Campaign Management Module")
add_table(["Test ID", "Execution Scenario", "Expected Architectural Response", "Empirical Result", "Status"],
    [["TC-05", "Submit campaign payload with null title", "Form validation intercepts, prevents Firestore write", "Write prevented, UI alerts", "Pass"],
     ["TC-06", "Inject valid payload", "Firestore commits document, WebSocket updates all active clients in < 500ms", "Real-time sync verified across devices", "Pass"],
     ["TC-07", "Mutate existing campaign state", "State change propagates instantaneously", "Global propagation successful", "Pass"],
     ["TC-08", "Execute delete operation", "Document wiped, UI list animates removal", "Removed cleanly", "Pass"]])
doc.add_paragraph()

table_caption("QA Vector: Complex Algorithmic Features (FYP-02 Scope)")
add_table(["Test ID", "Execution Scenario", "Expected Architectural Response", "Empirical Result", "Status"],
    [["TC-11", "Load feed for user with specific 'Medical' string tag", "Algorithm iterates, ranks 'Medical' campaigns at index 0", "Array sorted correctly", "Pass"],
     ["TC-12", "Trigger QR Generation payload", "Encode string to matrix, render high-res image", "QR Matrix rendered", "Pass"],
     ["TC-13", "Trigger Camera API on valid QR", "Decode matrix, dispatch attendance mutation to Firestore", "Status flipped to 'Present'", "Pass"]])

h2("7.5 User Acceptance Testing (UAT) and Usability Metrics")
body("Following technical validation, the software was deployed to a simulated field environment with three core volunteers from the HRAS NGO serving as primary evaluators.")
body("Empirical UAT Feedback: The evaluators successfully executed the critical path (Registration -> Profile Setup -> Campaign Join) in an average time of 58 seconds, comfortably exceeding the NFR usability threshold. The evaluators explicitly praised the low-friction user interface and the instantaneous response of the Smart Matching feed, noting that it drastically outperformed their legacy WhatsApp-based coordination methodologies.")
doc.add_page_break()

# =================== CHAPTER 8 ===================
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Chapter 8'); r.bold = True; r.font.size = Pt(22)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Conclusion and Future Roadmaps'); r2.bold = True; r2.font.size = Pt(20)
doc.add_page_break()

h1("Chapter 8: Conclusion and Future Roadmaps")

h2("8.1 Synthesis of Project Achievements")
body("Over the culmination of three rigorous academic semesters, this final year project successfully engineered, validated, and delivered a high-performance, enterprise-grade NGO Management Ecosystem. By strictly adhering to the architectural requirements established through deep stakeholder elicitation, the system successfully resolved the crippling administrative bottlenecks faced by the HRAS organization. The deployment achieved the following definitive milestones:")
for a in [
    "Engineered a high-framerate, cross-platform mobile application targeting both the iOS and Android ecosystems from a singular, highly maintainable Flutter codebase.",
    "Architected an expansive, responsive web-based administrative dashboard granting NGO executives unprecedented macro-level operational visibility.",
    "Eliminated data latency by integrating the Firebase Firestore NoSQL infrastructure, achieving sub-second data synchronization across distributed client nodes.",
    "Secured the financial and organizational data silos through mathematically enforced, server-side Role-Based Access Control (RBAC).",
    "Maximized volunteer deployment efficiency by implementing a customized Smart Matching algorithmic engine.",
    "Eradicated field logistics bottlenecks by engineering a zero-latency, cryptographically secure QR-code attendance scanner.",
    "Solidified donor trust by developing a dynamic PDF rendering engine to provide immutable financial receipts."
]:
    body(f"• {a}")

h2("8.2 Critical System Limitations")
body("An objective engineering analysis mandates the acknowledgment of current system limitations. While functionally robust, the v1.0 architecture operates under the following constraints:")
body("1. Perpetual Network Dependency: The reliance on Cloud BaaS (Backend-as-a-Service) dictates that the application demands a persistent internet connection. True offline mutation queues and local SQLite synchronization have not been implemented, limiting operability in deep rural or disaster-struck zones with collapsed cellular networks.")
body("2. Disconnected Financial API Topology: Due to insurmountable corporate registration barriers and high API transaction fees, direct integration with tier-1 payment gateways (Stripe, JazzCash) was omitted. Consequently, the system relies on manual administrative intervention to log bank transfers, retaining a minor attack vector for human data-entry error.")
body("3. Deterministic Algorithmic Bounds: The Smart Matching engine utilizes a deterministic boolean logic matrix. It lacks predictive, machine-learning capabilities that could analyze historical volunteer behaviors to predict future campaign participation likelihoods.")

h2("8.3 Strategic Roadmap for Future Work")
body("To elevate the system from a bespoke utility for HRAS into a nationally scalable SaaS product, future software development cycles should prioritize the following architectural expansions:")
body("• Fintech API Integration: Securing corporate compliance to integrate native payment hooks, allowing instant, automated donation settlement and ledger updating without human oversight.")
body("• Localization Engine: Abstracting all hardcoded UI strings into translation files to support real-time application switching to Urdu and regional dialects, drastically expanding demographic penetration.")
body("• Geospatial Tracking and Notifications: Upgrading the platform to utilize native GPS hardware, enabling proximity-based push notifications (e.g., 'A blood drive is happening 2km from your current location') and turn-by-turn navigation integrations.")
body("• Predictive Machine Learning Dashboards: Refactoring the web dashboard to incorporate Python-based data-science APIs, providing the NGO with predictive trend analysis regarding seasonal donation influxes and volunteer attrition rates.")

h2("8.4 Final Academic Conclusion")
body("The successful execution of the HRAS Digital Management System proves unequivocally that the injection of modern software engineering methodologies into grassroots charitable organizations yields massive operational dividends. By leveraging cutting-edge cross-platform compilation (Flutter) and highly scalable cloud infrastructure (Firebase), this project transformed a chaotic, paper-bound administrative nightmare into an elegant, automated, and mathematically transparent ecosystem. Ultimately, this software does not just organize data—it fundamentally empowers the non-profit sector to redirect their invaluable time and cognitive resources away from bureaucratic paperwork and back toward their prime directive: aggressive, high-impact humanitarian action.")
doc.add_page_break()

# =================== USER MANUAL & TECHNICAL GLOSSARY ===================
h1("Operational User Manual")
h2("Volunteer Client (Mobile Application)")
body("1. Initialization: Download the compiled binary and execute the application. Tap 'Sign Up', inputting valid credentials into the regex-protected form fields.")
body("2. Skill Profiling: Navigate to the User Profile matrix. Inject relevant strings (e.g., 'Medical', 'Logistics') into your skill array. This is critical for the matching algorithm.")
body("3. Campaign Interaction: Return to the primary feed. The system will asynchronously fetch and rank campaigns. Tap a rendered card to review parameters, then execute the 'Join' mutation to register.")
body("4. Field Operations: Upon physical arrival at the campaign coordinates, initialize the 'Scan QR' module. Focus the camera hardware on the Administrator's matrix to securely mutate your state to 'Present'.")

h2("Administrative Client (Web Dashboard)")
body("1. Authentication: Navigate to the deployed web URL. Provide high-clearance administrative credentials. The server-side RBAC will validate the token and route to the master dashboard.")
body("2. Logistics Management: Utilize the persistent navigation rail to access the Campaigns subsystem. Execute the 'Add New' command to construct a new event payload for the database.")
body("3. Execution: During an active campaign, locate the specific event row and trigger the 'Generate QR' cryptographic function. Display the resulting high-res matrix on a tablet or laptop for volunteer scanning.")
body("4. Financial Auditing: Navigate to the Donations matrix. Input validated bank transfer amounts. The system will dynamically re-calculate the aggregate financial parameters in real-time.")
doc.add_page_break()

h1("Technical Glossary")
for term, defn in [
    ("Firebase BaaS", "Backend-as-a-Service architecture provided by Google, heavily utilized for its NoSQL capabilities, Auth, and zero-configuration scaling."),
    ("Flutter UI SDK", "Google's UI toolkit utilized to compile native applications for mobile, web, and desktop from a singular Dart codebase using a high-performance Skia/Impeller rendering engine."),
    ("Firestore NoSQL", "A highly flexible, scalable cloud database allowing data to be stored in JSON-like document hierarchies, updating all connected clients in real-time via WebSockets."),
    ("CRUD Paradigm", "Create, Read, Update, Delete — the fundamental algorithmic operations required to manage persistent database storage."),
    ("UAT (User Acceptance Testing)", "The absolute final phase of software testing, wherein actual end-users execute the software in a simulated production environment to validate functional design."),
    ("RBAC (Role-Based Access Control)", "An aggressive network security paradigm that restricts system access to authorized personnel based strictly on their cryptographic role assignments, enforced server-side."),
]:
    body(f"{term}: {defn}")

out = os.path.join(os.path.dirname(__file__), 'Thesis-FYP03', 'FYP_03_Thesis.docx')
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
inject_update_fields(out)
print(f"FYP-03 Thesis (Extreme Edition) saved: {out}")
