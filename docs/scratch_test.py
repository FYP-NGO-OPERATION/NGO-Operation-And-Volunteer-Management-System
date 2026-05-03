import docx
from docx.oxml import OxmlElement, ns
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_seq_field(paragraph, seq_identifier, text_before, text_after):
    """
    Adds a sequence field to a paragraph.
    seq_identifier: 'Figure' or 'Table'
    """
    run_before = paragraph.add_run(text_before)
    run_before.font.size = Pt(10)
    
    # <w:fldChar w:fldCharType="begin"/>
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(ns.qn('w:fldCharType'), 'begin')
    r1 = paragraph.add_run()
    r1._r.append(fldChar1)
    
    # <w:instrText xml:space="preserve"> SEQ Figure \* ARABIC </w:instrText>
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = f' SEQ {seq_identifier} \\* ARABIC '
    r2 = paragraph.add_run()
    r2._r.append(instrText)
    
    # <w:fldChar w:fldCharType="separate"/>
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')
    r3 = paragraph.add_run()
    r3._r.append(fldChar2)
    
    # Placeholder number (Word will update this)
    r4 = paragraph.add_run('1')
    r4.font.size = Pt(10)
    
    # <w:fldChar w:fldCharType="end"/>
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(ns.qn('w:fldCharType'), 'end')
    r5 = paragraph.add_run()
    r5._r.append(fldChar3)
    
    run_after = paragraph.add_run(text_after)
    run_after.font.size = Pt(10)

def add_toc_field(doc, title, switch):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(16); r.font.name = 'Times New Roman'
    doc.add_paragraph()
    
    p2 = doc.add_paragraph()
    run = p2.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(ns.qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = f' TOC {switch} '
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(ns.qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

doc = docx.Document()
add_toc_field(doc, "Table of Contents", "\\o \"1-3\" \\h \\z \\u")
doc.add_page_break()
add_toc_field(doc, "List of Tables", "\\h \\z \\c \"Table\"")
doc.add_page_break()
add_toc_field(doc, "List of Figures", "\\h \\z \\c \"Figure\"")
doc.add_page_break()

doc.add_heading("Chapter 1", level=1)

p_table = doc.add_paragraph()
p_table.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_seq_field(p_table, "Table", "Table ", ": My First Table")

doc.add_paragraph("Table content here")

p_fig = doc.add_paragraph()
p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_seq_field(p_fig, "Figure", "Figure ", ": My First Figure")

doc.save("test_fields.docx")
print("Saved test_fields.docx")
