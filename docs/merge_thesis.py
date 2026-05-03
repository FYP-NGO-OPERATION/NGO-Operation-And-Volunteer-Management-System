"""Merge FYP-01/02/03 thesis DOCX files into one FINAL_THESIS_COMPLETE.docx."""
import os, copy, zipfile, tempfile, shutil
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.enum.section import WD_SECTION

BASE = os.path.dirname(__file__)

def inject_update_fields(docx_path):
    """Unzips the docx, injects updateFields into settings.xml, and rezips it."""
    temp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
            
        settings_path = os.path.join(temp_dir, 'word', 'settings.xml')
        if os.path.exists(settings_path):
            with open(settings_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Inject w:updateFields if not present
            if '<w:updateFields w:val="true"/>' not in content:
                # Find <w:settings ...> and insert right after
                import re
                content = re.sub(r'(<w:settings[^>]*>)', r'\1<w:updateFields w:val="true"/>', content, count=1)
                
                with open(settings_path, 'w', encoding='utf-8') as f:
                    f.write(content)
        
        # Rezip
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zip_out.write(file_path, arcname)
    finally:
        shutil.rmtree(temp_dir)

def merge():
    # Use FYP-01 as the base (has front matter + Ch 1-2)
    base_path = os.path.join(BASE, 'Thesis-FYP01', 'FYP_01_Final_Submission.docx')
    doc = Document(base_path)

    # Load FYP-02 (Ch 3-6) and FYP-03 (Ch 7-8)
    fyp02 = Document(os.path.join(BASE, 'Thesis-FYP02', 'FYP_02_Thesis.docx'))
    fyp03 = Document(os.path.join(BASE, 'Thesis-FYP03', 'FYP_03_Thesis.docx'))

    def append_doc(target, source, skip_title_page=True):
        """Append paragraphs and tables from source to target."""
        started = not skip_title_page
        for element in source.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
            if tag == 'sectPr':
                continue
            # Skip title page (first few elements until we hit a Heading 1 with "Chapter")
            if not started:
                if tag == 'p':
                    text = element.text or ''
                    if 'Chapter 3' in text or 'Chapter 7' in text:
                        started = True
                    else:
                        continue
                else:
                    continue
            # Copy element over
            copied = copy.deepcopy(element)
            target.element.body.append(copied)

    # Add page break before Ch 3
    doc.add_page_break()
    append_doc(doc, fyp02, skip_title_page=True)

    # Add page break before Ch 7
    doc.add_page_break()
    append_doc(doc, fyp03, skip_title_page=True)

    # Save initial
    out = os.path.join(BASE, 'FINAL_THESIS_COMPLETE.docx')
    doc.save(out)
    
    # Inject auto-update tag
    inject_update_fields(out)
    
    print(f"FINAL THESIS saved: {out}")
    print("Auto-update fields flag injected. Word will prompt to update fields on open.")

    # Verify structure
    headings = []
    final = Document(out)
    for p in final.paragraphs:
        if p.style.name.startswith('Heading 1'):
            headings.append(p.text)
    print(f"\nHeading 1 entries found: {len(headings)}")
    for h in headings:
        print(f"  - {h}")

if __name__ == '__main__':
    merge()
